import datetime as dt
import os
import re

from docx import Document
from docx.enum.table import (WD_ALIGN_VERTICAL, WD_CELL_VERTICAL_ALIGNMENT,
                             WD_TABLE_ALIGNMENT)
from docx.enum.text import (WD_ALIGN_PARAGRAPH, WD_LINE_SPACING,
                            WD_PARAGRAPH_ALIGNMENT)
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

xmwz=str(input('请输入项目位置：'))
xmbh=str(input('请输入项目编号：'))
path=os.getcwd()
# os.mkdir('02计算数据')
a=os.listdir(path)
path_list=path.split('\\')
xmmc=path_list[-2]
xmmc=re.sub(r'[GC]+\d+','',xmmc)
new_name='RL01'+xmmc+'竣工测量报告'
for i in a:
    if 'txt' in i:
        name=i
ywtj=[]
ywzcd=[]
ewtj=[]
ewzcd=[]
zongtj=[]
with open(name,'rt') as gzl:
    gcllist=gzl.readlines()
    index=0
    for item in gcllist:
        if '一次网总长度'  in item :
            ywwz=index
            line=str(item).replace('\n','')
            ywzcd.append(line.split('：',1))
        elif '二次网总长度' in item :
            ewwz=index
            line=str(item).replace('\n','')
            ewzcd.append(line.split('：',1))
        index+=1
    yiwanglist=gcllist[1:ewwz]
    erwanglist=gcllist[ewwz+1:len(gcllist)-1]
    zonglist=gcllist[len(gcllist)-1:len(gcllist)]
    for  yiwang in yiwanglist:
        line=str(yiwang).replace('\n','')
        ywtj.append(line.split('：',1))
    for erwang in erwanglist:
        line=str(erwang).replace('\n','')
        ewtj.append(line.split('：',1))
    for zong in zonglist:
        line=str(zong).replace('\n','')
        zongtj.append(line.split('：',1))

document=Document('C:/模板/普天热力报告模板.docx')
document.add_page_break()    #添加分页符
if ewwz>0:
    p = document.add_paragraph()
    p.paragraph_format.space_after = Pt(0)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
    p.paragraph_format.line_spacing = 2.0
    p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run=p.add_run('管网总长度统计表')
    run.font.size=Pt(12)   #字体
    run.font.name='宋体'
    run.element.rPr.rFonts.set(qn('w:eastAsia'),'宋体')
    Table=document.add_table(rows= 4,cols= 4,style= 'Table Grid')
    Table.alignment = WD_TABLE_ALIGNMENT.CENTER # 设置表格为居中对齐
    for i in range(4):
        Table.rows[i].height=Cm(0.7)
        for col in range(4):
            Table.cell(i,col).vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            Table.cell(i,col).paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            Table.cell(i,col).width = Cm(2.5)
    Table.cell(0,0).paragraphs[0].text='序  号'
    Table.cell(0,1).paragraphs[0].text='类  型'
    Table.cell(0,2).paragraphs[0].text='长度（米）'
    Table.cell(0,3).paragraphs[0].text='备  注'
    Table.cell(1,0).paragraphs[0].text='1'
    Table.cell(1,1).paragraphs[0].text='一次网'
    Table.cell(1,2).paragraphs[0].text=ywzcd[0][1]
    Table.cell(1,3).paragraphs[0].text=''
    Table.cell(2,0).paragraphs[0].text='2'
    Table.cell(2,1).paragraphs[0].text='二次网'
    Table.cell(2,2).paragraphs[0].text=ewzcd[0][1]
    Table.cell(2,3).paragraphs[0].text=''
    Table.cell(3,0).paragraphs[0].text='3'
    Table.cell(3,1).paragraphs[0].text='总计'
    Table.cell(3,2).paragraphs[0].text=zongtj[0][1]
    Table.cell(3,3).paragraphs[0].text=''  
    p = document.add_paragraph()
    p.paragraph_format.space_after = Pt(0)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
    p.paragraph_format.line_spacing = 3.0
    p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run=p.add_run('一次网长度统计表')
    run.font.size=Pt(12)   #字体
    run.font.name='宋体'
    run.element.rPr.rFonts.set(qn('w:eastAsia'),'宋体')
    Table1=document.add_table(rows= len(ywtj)+2,cols= 4,style= 'Table Grid')
    Table1.alignment = WD_TABLE_ALIGNMENT.CENTER # 设置表格为居中对齐
    for i in range(len(ywtj)+2):
        Table1.rows[i].height=Cm(0.7)
        for col in range(4):
            Table1.cell(i,col).vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            Table1.cell(i,col).paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            Table1.cell(i,col).width = Cm(2.5)
    Table1.cell(0,0).paragraphs[0].text='序  号'
    Table1.cell(0,1).paragraphs[0].text='类  型'
    Table1.cell(0,2).paragraphs[0].text='长度（米）'
    Table1.cell(0,3).paragraphs[0].text='备  注'
    for i in range(len(yiwanglist)):
        Table1.cell(i+1,0).paragraphs[0].text=str(i+1)
        Table1.cell(i+1,1).paragraphs[0].text=ywtj[i][0]
        Table1.cell(i+1,2).paragraphs[0].text=ywtj[i][1]
    Table1.cell(len(yiwanglist)+1,0).paragraphs[0].text=str(len(yiwanglist)+1)
    Table1.cell(len(yiwanglist)+1,1).paragraphs[0].text='合计'
    Table1.cell(len(yiwanglist)+1,2).paragraphs[0].text=ywzcd[0][1]
    p = document.add_paragraph()
    p.paragraph_format.space_after = Pt(0)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
    p.paragraph_format.line_spacing = 3.0
    p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run=p.add_run('二次网长度统计表')
    run.font.size=Pt(12)   #字体
    run.font.name='宋体'
    run.element.rPr.rFonts.set(qn('w:eastAsia'),'宋体')
    Table2=document.add_table(rows= len(ewtj)+2,cols= 4,style= 'Table Grid')
    Table2.alignment = WD_TABLE_ALIGNMENT.CENTER # 设置表格为居中对齐
    for i in range(len(ewtj)+2):
        Table2.rows[i].height=Cm(0.7)
        for col in range(4):
            Table2.cell(i,col).vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            Table2.cell(i,col).paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            Table2.cell(i,col).width = Cm(2.5)
    Table2.cell(0,0).paragraphs[0].text='序  号'
    Table2.cell(0,1).paragraphs[0].text='类  型'
    Table2.cell(0,2).paragraphs[0].text='长度（米）'
    Table2.cell(0,3).paragraphs[0].text='备  注'
    for i in range(len(erwanglist)):
        Table2.cell(i+1,0).paragraphs[0].text=str(i+1)
        Table2.cell(i+1,1).paragraphs[0].text=ewtj[i][0]
        Table2.cell(i+1,2).paragraphs[0].text=ewtj[i][1]
    Table2.cell(len(erwanglist)+1,0).paragraphs[0].text=str(len(erwanglist)+1)
    Table2.cell(len(erwanglist)+1,1).paragraphs[0].text='合计'
    Table2.cell(len(erwanglist)+1,2).paragraphs[0].text=ewzcd[0][1]
if ewwz==0:
    p = document.add_paragraph()
    p.paragraph_format.space_after = Pt(0)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run=p.add_run('二次网长度统计表')
    run.font.size=Pt(12)   #字体
    run.font.name='宋体'
    run.element.rPr.rFonts.set(qn('w:eastAsia'),'宋体')
    Table2=document.add_table(rows= len(ewtj)+2,cols= 4,style= 'Table Grid')
    Table2.alignment = WD_TABLE_ALIGNMENT.CENTER # 设置表格为居中对齐
    for i in range(len(ewtj)+2):
        Table2.rows[i].height=Cm(0.7)
        for col in range(4):
            Table2.cell(i,col).vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            Table2.cell(i,col).paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            Table2.cell(i,col).width = Cm(2.5)
    Table2.cell(0,0).paragraphs[0].text='序  号'
    Table2.cell(0,1).paragraphs[0].text='类  型'
    Table2.cell(0,2).paragraphs[0].text='长度（米）'
    Table2.cell(0,3).paragraphs[0].text='备  注'
    for i in range(len(erwanglist)):
        Table2.cell(i+1,0).paragraphs[0].text=str(i+1)
        Table2.cell(i+1,1).paragraphs[0].text=ewtj[i][0]
        Table2.cell(i+1,2).paragraphs[0].text=ewtj[i][1]
    Table2.cell(len(erwanglist)+1,0).paragraphs[0].text=str(len(erwanglist)+1)
    Table2.cell(len(erwanglist)+1,1).paragraphs[0].text='合计'
    Table2.cell(len(erwanglist)+1,2).paragraphs[0].text=ewzcd[0][1]


for paragraph in document.paragraphs:
        for run in paragraph.runs:
            if '项目名称' in run.text:
                run.text=run.text.replace('项目名称',xmmc)
            elif '工程量' in run.text:
                run.text=run.text.replace('工程量',zongtj[0][1])
            elif '项目地址' in run.text:
                run.text=run.text.replace('项目地址',xmwz)
            elif '报告日期' in run.text:
                date=dt.datetime.now()
                date_now=str(date.year)+'年'+str(date.month)+'月'
                run.text=run.text.replace('报告日期',date_now)
            elif '填写编号' in run.text:
                run.text=run.text.replace('填写编号',xmbh)
            elif '测量日期' in run.text:
                run.text=run.text.replace('测量日期',date_now)
document.save(new_name+'.docx')
# document.save(new_name+'.docx')                
# def read_document():  
#     for paragraph in document.paragraphs:
#         for run in paragraph.runs:
#             if '项目名称' in run.text:
#                 run.text=run.text.replace('项目名称','工程量')
#             elif '工程量' in run.text:
#                 run.text=run.text.replace('工程量','工程量')
#             elif '项目地址' in run.text:
#                 run.text=run.text.replace('项目地址','工程量')
#             elif '报告日期' in run.text:
#                 date=dt.datetime.now()
#                 date_now=str(date.year)+'年'+str(date.month)+'月'
#                 run.text=run.text.replace('报告日期','工程量')
#             elif '填写编号' in run.text:
#                 run.text=run.text.replace('填写编号','工程量')
#             elif '测量日期' in run.text:
#                 run.text=run.text.replace('测量日期','工程量')
#     document.save('new_name'+'.docx')